import os
import pytest
from unittest.mock import patch, MagicMock
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from datetime import date, datetime

# Set up test database (file-based SQLite for persistence across requests)
from app.database import Base, get_db
from app.main import app
from app.models.user import User
from app.models.company import Company
from app.models.vehicle import Vehicle
from app.models.bill import Bill
from app.models.payment import Payment
from app.models.audit_log import AuditLog
from app.utils.security import hash_password, create_access_token
from app.services.docx_segmenter import BillChunk
from app.services.document_intelligence import DocumentIntelligenceService


TEST_DB_FILE = "./test_temp.db"
SQLALCHEMY_DATABASE_URL = f"sqlite:///{TEST_DB_FILE}"

engine = create_engine(
    SQLALCHEMY_DATABASE_URL, 
    connect_args={"check_same_thread": False}
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# Override database dependency
def override_get_db():
    db = TestingSessionLocal()
    try:
        yield db
    finally:
        db.close()

app.dependency_overrides[get_db] = override_get_db

client = TestClient(app)

@pytest.fixture(scope="module", autouse=True)
def setup_database():
    # Remove existing temp DB if it somehow exists
    if os.path.exists(TEST_DB_FILE):
        os.remove(TEST_DB_FILE)
        
    Base.metadata.create_all(bind=engine)
    db = TestingSessionLocal()
    
    # Create test users with different roles
    owner = User(
        username="owner_test",
        password=hash_password("admin123"),
        email="owner@test.com",
        role="OWNER",
        active=True
    )
    manager = User(
        username="manager_test",
        password=hash_password("manager123"),
        email="manager@test.com",
        role="MANAGER",
        active=True
    )
    employee = User(
        username="employee_test",
        password=hash_password("employee123"),
        email="employee@test.com",
        role="EMPLOYEE",
        active=True
    )
    
    db.add(owner)
    db.add(manager)
    db.add(employee)
    db.commit()
    db.close()
    
    yield
    
    # Teardown: drop metadata and delete database file
    Base.metadata.drop_all(bind=engine)
    if os.path.exists(TEST_DB_FILE):
        # We try to remove but ignore error if file is locked on Windows temporarily
        try:
            os.remove(TEST_DB_FILE)
        except Exception:
            pass

# Helper function to generate auth headers
def get_auth_headers(username: str, role: str):
    token = create_access_token(username, role)
    return {"Authorization": f"Bearer {token}"}

# ==========================================
# AUTHENTICATION & SECURITY TESTS
# ==========================================

def test_login_success():
    payload = {"username": "owner_test", "password": "admin123"}
    response = client.post("/api/auth/login", json=payload)
    assert response.status_code == 200
    json_data = response.json()
    assert "token" in json_data
    assert json_data["role"] == "OWNER"

def test_login_invalid_credentials():
    payload = {"username": "owner_test", "password": "wrongpassword"}
    response = client.post("/api/auth/login", json=payload)
    assert response.status_code == 401
    assert response.json()["detail"] == "Invalid username or password"

def test_register_role_restriction():
    # Employee attempts to register another user (should fail)
    headers = get_auth_headers("employee_test", "EMPLOYEE")
    payload = {
        "username": "new_manager",
        "password": "password123",
        "email": "new_man@test.com",
        "role": "MANAGER"
    }
    response = client.post("/api/auth/register", json=payload, headers=headers)
    assert response.status_code == 403

def test_register_by_owner_success():
    headers = get_auth_headers("owner_test", "OWNER")
    payload = {
        "username": "new_manager",
        "password": "password123",
        "email": "new_man@test.com",
        "role": "MANAGER"
    }
    response = client.post("/api/auth/register", json=payload, headers=headers)
    assert response.status_code == 201
    assert response.json()["username"] == "new_manager"

# ==========================================
# COMPANIES CRUD TESTS
# ==========================================

def test_company_crud():
    owner_headers = get_auth_headers("owner_test", "OWNER")
    employee_headers = get_auth_headers("employee_test", "EMPLOYEE")
    
    # 1. Create company (Owner success)
    company_payload = {
        "name": "Test Company Corp",
        "address": "123 Test St, Hyderabad",
        "gstNumber": "36AAAAA0000A1Z1"
    }
    response = client.post("/api/companies", json=company_payload, headers=owner_headers)
    assert response.status_code == 200
    company_id = response.json()["id"]
    assert response.json()["name"] == "Test Company Corp"

    # 2. Create company (Employee permission check: Manager/Owner only)
    response = client.post("/api/companies", json=company_payload, headers=employee_headers)
    assert response.status_code == 403

    # 3. Read companies (Any role success)
    response = client.get("/api/companies", headers=employee_headers)
    assert response.status_code == 200
    assert len(response.json()) >= 1

    # 4. Update company (Owner success)
    update_payload = {
        "name": "Test Company Corp Updated",
        "address": "456 Main Rd, Hyderabad",
        "gstNumber": "36AAAAA0000A2Z2"
    }
    response = client.put(f"/api/companies/{company_id}", json=update_payload, headers=owner_headers)
    assert response.status_code == 200
    assert response.json()["name"] == "Test Company Corp Updated"

    # 5. Delete company (Owner success)
    response = client.delete(f"/api/companies/{company_id}", headers=owner_headers)
    assert response.status_code == 204

# ==========================================
# VEHICLES CRUD TESTS
# ==========================================

def test_vehicle_crud():
    owner_headers = get_auth_headers("owner_test", "OWNER")
    
    # 1. Create vehicle
    vehicle_payload = {
        "registrationNumber": "AP-09-TV-1234",
        "type": "SUV",
        "model": "Innova Crysta"
    }
    response = client.post("/api/vehicles", json=vehicle_payload, headers=owner_headers)
    assert response.status_code == 200
    vehicle_id = response.json()["id"]
    assert response.json()["registrationNumber"] == "AP-09-TV-1234"

    # 2. Get vehicles list
    response = client.get("/api/vehicles", headers=owner_headers)
    assert response.status_code == 200
    assert len(response.json()) >= 1

    # 3. Update vehicle
    update_payload = {
        "registrationNumber": "AP-09-TV-1234",
        "type": "SUV",
        "model": "Innova Crysta Luxury"
    }
    response = client.put(f"/api/vehicles/{vehicle_id}", json=update_payload, headers=owner_headers)
    assert response.status_code == 200
    assert response.json()["model"] == "Innova Crysta Luxury"

    # 4. Delete vehicle
    response = client.delete(f"/api/vehicles/{vehicle_id}", headers=owner_headers)
    assert response.status_code == 204

# ==========================================
# BILLS CRUD & CALCULATIONS TESTS
# ==========================================

def test_bill_crud():
    owner_headers = get_auth_headers("owner_test", "OWNER")
    
    # Pre-create company and vehicle for integrity constraints
    client.post("/api/companies", json={"name": "Ashapura Travels", "address": "Hyd", "gstNumber": "36AA"}, headers=owner_headers)
    client.post("/api/vehicles", json={"registrationNumber": "AP-10-XX-9999", "type": "Sedan", "model": "Dezire"}, headers=owner_headers)

    bill_payload = {
        "billDate": "2026-06-30",
        "companyName": "Ashapura Travels",
        "vehicleName": "AP-10-XX-9999",
        "dutySlipNo": "DS-9041",
        "tripDate": "2026-06-29",
        "vehicleType": "Sedan",
        "acNonAc": "AC",
        "totalKms": 300.0,
        "totalHours": 12.0,
        "tripType": "Local",
        "pricingType": "BASE",
        "baseAmount": 2500.0,
        "driverBata": 300.0,
        "parking": 150.0,
        "toll": 200.0,
        "nightCharges": 100.0,
        "otherCharges": 50.0,
        "notes": "Test trip notes details",
        "contactPerson": "Ramesh Kumar",
        "bookedBy": "Owner Admin",
        "managerName": "Suresh Manager"
    }

    # 1. Create bill & verify grand total calculation (sum of base + bata + toll etc)
    # Expected grand total = 2500 + 300 + 150 + 200 + 100 + 50 = 3300.00
    response = client.post("/api/bills", json=bill_payload, headers=owner_headers)
    assert response.status_code == 201
    bill_data = response.json()
    bill_id = bill_data["id"]
    assert bill_data["grandTotal"] == 3300.0
    assert bill_data["billNumber"].startswith("BILL-")

    # 2. Get list of bills
    response = client.get("/api/bills", headers=owner_headers)
    assert response.status_code == 200
    assert "content" in response.json()
    assert len(response.json()["content"]) >= 1

    # 3. Get single bill details
    response = client.get(f"/api/bills/{bill_id}", headers=owner_headers)
    assert response.status_code == 200
    assert response.json()["dutySlipNo"] == "DS-9041"

    # 4. Search bills
    response = client.get("/api/bills/search?companyName=Ashapura", headers=owner_headers)
    assert response.status_code == 200
    assert len(response.json()["content"]) >= 1

    # 5. Export PDF invoice check
    response = client.get(f"/api/bills/{bill_id}/pdf", headers=owner_headers)
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert len(response.content) > 100  # Should return PDF stream bytes

    # 6. Delete bill (Owner success, Manager/Employee fails)
    response = client.delete(f"/api/bills/{bill_id}", headers=owner_headers)
    assert response.status_code == 204

# ==========================================
# AI SERVICE ROUTER PATCHED TESTS
# ==========================================

@patch("app.services.gemini.GeminiService._post")
def test_ai_insights_endpoint(mock_post):
    owner_headers = get_auth_headers("owner_test", "OWNER")
    
    # Mock AI Service success payload
    mock_post.return_value.status_code = 200
    mock_post.return_value.json.return_value = {
        "insights": [
            {"type": "TREND", "message": "Ashapura revenue is up by 15%", "confidence": 0.95}
        ]
    }

    response = client.get("/api/analytics/ai-insights", headers=owner_headers)
    assert response.status_code == 200
    assert response.json()["insights"][0]["type"] == "TREND"

@patch("app.services.gemini.GeminiService._post")
def test_ai_search_nl_explain(mock_post):
    owner_headers = get_auth_headers("owner_test", "OWNER")
    
    # Mock AI response translating NL to JSON filters
    mock_post.return_value.status_code = 200
    mock_post.return_value.json.return_value = {
        "companyName": "Ashapura Travels",
        "minAmount": 5000.0,
        "summary": "Ashapura bills above 5000"
    }

    response = client.get("/api/bills/search/nl/explain?query=Ashapura bills above 5000", headers=owner_headers)
    assert response.status_code == 200
    assert response.json()["companyName"] == "Ashapura Travels"
    assert response.json()["minAmount"] == 5000.0

@patch("app.services.imports.DocxSegmenterService.segment_docx")
@patch("app.services.imports.AiExtractionService.extract_page_data")
@patch("app.services.imports.DocumentIntelligenceService.extract_document")
@patch("app.config.settings")
def test_ai_parse_endpoint(mock_settings, mock_doc_intel, mock_extract, mock_segment):
    owner_headers = get_auth_headers("owner_test", "OWNER")

    # Force feature flags so the test exercises the AiExtractionService path
    # (not the USE_ENTERPRISE_LABELER branch which bypasses extract_page_data)
    mock_settings.USE_ENTERPRISE_LABELER = False
    mock_settings.USE_ENTERPRISE_VALIDATION = False
    mock_settings.USE_ENTERPRISE_LEARNING = False

    # Mock a minimal valid DocumentIntelligenceService result
    mock_doc = MagicMock()
    mock_doc.to_json.return_value = {"metadata": {"file_name": "test.docx"}, "pages": [], "tables": [], "paragraphs": []}
    mock_doc.pages = []
    mock_doc_intel.return_value = mock_doc

    mock_segment.return_value = [
        BillChunk(
            page_number=1,
            company_name="Ashapura Travels",
            raw_text="raw extracted text",
            extracted_tables=[]
        )
    ]
    mock_extract.return_value = {
        "company": "Ashapura Travels",
        "billNumber": "DS-9041-TEST",
        "invoiceNumber": "DS-9041-TEST",
        "dutySlip": "DS-9041-TEST",
        "vehicleNumber": "AP-10-XX-9999",
        "vehicleType": "Sedan",
        "driver": "Ramesh",
        "reportingDate": "2026-06-30",
        "totalHours": 12.0,
        "totalKilometers": 300.0,
        "toll": 200.0,
        "parking": 150.0,
        "driverBata": 300.0,
        "totalAmount": 3300.0
    }

    files = {"files": ("test.docx", b"dummy file content", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    response = client.post("/api/import/ai-parse", files=files, headers=owner_headers)
    assert response.status_code == 200
    assert len(response.json()) == 1
    assert response.json()[0]["dutySlipNo"] == "DS-9041-TEST"
    assert "documentIntelligence" in response.json()[0]


def test_document_intelligence_service_parses_docx_structure():
    from io import BytesIO
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Header line")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Duty Slip"
    table.cell(0, 1).text = "DS-1001"
    table.cell(1, 0).text = "Vehicle Number"
    table.cell(1, 1).text = "AP-09-TV-1234"

    buffer = BytesIO()
    doc.save(buffer)

    parsed = DocumentIntelligenceService.extract_document(buffer.getvalue(), "sample.docx")
    payload = parsed.to_json()

    assert payload["metadata"]["file_name"] == "sample.docx"
    assert payload["pages"]
    assert payload["tables"]
    assert payload["tables"][0]["number_of_rows"] == 2
    assert payload["tables"][0]["number_of_columns"] == 2
    assert payload["paragraphs"][0]["text"] == "Header line"

def test_validation_service_checks():
    from app.services.business_validation_service import ValidationService
    from app.schemas.ai import AiBillResponse, AiBillCharge
    from sqlalchemy.orm import Session
    
    db = TestingSessionLocal()
    
    # Test valid bill (should return 0 warnings)
    valid_bill = AiBillResponse(
        dutySlipNo="DS-VAL-101",
        billDate="2026-06-30",
        companyName="Val Corp",
        vehicleNumber="AP-10-XY-1234",
        vehicleType="SUV",
        totalKms="100.0",
        totalHours="10.0",
        dynamicCharges=[
            AiBillCharge(name="Toll", amount="100.0"),
            AiBillCharge(name="Parking", amount="50.0")
        ],
        totalAmount="2500.0"
    )
    warnings = ValidationService.validate_bill(db, valid_bill)
    assert len(warnings) == 0
    
    # Test invalid bill with warnings
    invalid_bill = AiBillResponse(
        dutySlipNo="", # Missing
        billDate="30/06/2026", # Invalid format
        companyName="Val Corp",
        vehicleNumber="INVALID_PLATE", # Malformed
        vehicleType="SUV",
        dynamicCharges=[
            AiBillCharge(name="Toll", amount="3000.0") # Sum exceeds total amount
        ],
        totalAmount="1000.0"
    )
    warnings = ValidationService.validate_bill(db, invalid_bill)
    assert len(warnings) > 0
    assert any("Missing mandatory field" in w for w in warnings)
    assert any("Invalid date format" in w for w in warnings)
    assert any("Malformed vehicle" in w for w in warnings)
    assert any("Arithmetic warning" in w for w in warnings)
    db.close()


# ==========================================
# REPORTS, DASHBOARDS, BACKUPS & AUDIT TESTS
# ==========================================

def test_reports_endpoints():
    owner_headers = get_auth_headers("owner_test", "OWNER")
    
    # 1. Summary
    response = client.get("/api/reports/summary", headers=owner_headers)
    assert response.status_code == 200
    assert "totalBillsCount" in response.json()

    # 2. Top companies
    response = client.get("/api/reports/top-companies", headers=owner_headers)
    assert response.status_code == 200
    assert isinstance(response.json(), list)

    # 3. Top vehicles
    response = client.get("/api/reports/top-vehicles", headers=owner_headers)
    assert response.status_code == 200
    assert isinstance(response.json(), list)

def test_dashboard_endpoint():
    owner_headers = get_auth_headers("owner_test", "OWNER")
    
    response = client.get("/api/dashboard/owner", headers=owner_headers)
    assert response.status_code == 200
    assert "stats" in response.json()
    assert "revenueTrend" in response.json()

def test_audit_logs_logging():
    owner_headers = get_auth_headers("owner_test", "OWNER")
    
    response = client.get("/api/audit-logs", headers=owner_headers)
    assert response.status_code == 200
    data = response.json()
    assert isinstance(data, dict)
    assert "content" in data
    assert isinstance(data["content"], list)
    # Ensure standard audit actions are mapped
    assert len(data["content"]) >= 1




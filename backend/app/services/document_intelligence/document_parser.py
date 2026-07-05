from __future__ import annotations

import io
import logging
from typing import List, Optional

from docx import Document

from app.services.docx_extractor import DocxExtractionService
from app.services.document_intelligence.coordinate_mapper import build_source_path
from app.services.document_intelligence.document_models import (
    DocumentCell,
    DocumentMetadata,
    DocumentPage,
    DocumentParagraph,
    DocumentRun,
    DocumentLine,
    DocumentTable,
    EnterpriseDocument,
)
from app.services.document_intelligence.metadata_parser import extract_document_metadata
from app.services.document_intelligence.page_parser import parse_pages

logger = logging.getLogger("document_intelligence")


class DocumentIntelligenceService:
    @staticmethod
    def extract_document(file_bytes: bytes, file_name: Optional[str] = None, mime_type: Optional[str] = None) -> EnterpriseDocument:
        """
        Build a structured, deterministic representation of the Word document.
        The service preserves document structure and falls back gracefully on malformed input.
        """
        file_name = file_name or "unknown"
        errors: List[str] = []
        warnings: List[str] = []

        try:
            document = Document(io.BytesIO(file_bytes))
            metadata_dict = extract_document_metadata(document, file_name, len(file_bytes), mime_type)
            pages, reading_order, coordinates, page_warnings = parse_pages(document, file_name)
            warnings.extend(page_warnings)

            paragraphs: List[DocumentParagraph] = []
            tables: List[DocumentTable] = []
            cells: List[DocumentCell] = []
            lines: List[DocumentLine] = []
            runs: List[DocumentRun] = []

            for page in pages:
                paragraphs.extend(page.paragraphs)
                tables.extend(page.tables)
                for paragraph in page.paragraphs:
                    lines.extend(paragraph.lines)
                    runs.extend(paragraph.runs)
                for table in page.tables:
                    for row in table.rows:
                        cells.extend(row.cells)

            document_model = EnterpriseDocument(
                id=build_source_path("document", file_name, len(file_bytes)),
                metadata=DocumentMetadata(**metadata_dict, warnings=warnings, errors=errors),
                pages=pages,
                paragraphs=paragraphs,
                tables=tables,
                cells=cells,
                lines=lines,
                runs=runs,
                coordinates=coordinates,
                reading_order=reading_order,
                errors=errors,
                warnings=warnings,
            )
            return document_model
        except Exception as exc:
            logger.warning("Document intelligence parse failed for %s: %s", file_name, exc)
            errors.append(str(exc))

            fallback_text = DocxExtractionService.extract_raw_text(file_bytes, file_name)
            fallback_metadata = DocumentMetadata(
                file_name=file_name,
                file_size_bytes=len(file_bytes),
                mime_type=mime_type,
                source_format="legacy_or_malformed",
                warnings=["Document parsed via fallback text decoder."],
                errors=errors,
            )
            fallback_page = DocumentPage(
                id=build_source_path("document", file_name, "fallback"),
                page_number=1,
                position=1,
                title=file_name,
                paragraphs=[],
                tables=[],
                reading_order=[],
                metadata={"fallback_text": fallback_text},
            )
            return EnterpriseDocument(
                id=build_source_path("document", file_name, "fallback", len(file_bytes)),
                metadata=fallback_metadata,
                pages=[fallback_page],
                coordinates=[],
                reading_order=[],
                errors=errors,
                warnings=["Malformed DOCX handled gracefully."],
            )

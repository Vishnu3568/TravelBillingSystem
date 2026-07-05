from __future__ import annotations

from typing import Any, Dict, List, Tuple

from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.document_intelligence.coordinate_mapper import build_source_path, make_node_id, make_coordinate_id
from app.services.document_intelligence.document_models import DocumentCoordinate, DocumentPage
from app.services.document_intelligence.layout_parser import build_paragraph_models
from app.services.document_intelligence.table_parser import parse_table


def _has_page_break_before(child) -> bool:
    p_pr = child.find(qn("w:pPr"))
    if p_pr is None:
        return False
    page_break_before = p_pr.find(qn("w:pageBreakBefore"))
    return page_break_before is not None


def _has_page_break_within(child) -> bool:
    if child.xpath('.//w:br[@w:type="page"]'):
        return True
    if child.xpath(".//w:lastRenderedPageBreak"):
        return True
    return False


def parse_pages(doc, file_name: str) -> Tuple[List[DocumentPage], List[Dict[str, Any]], List[DocumentCoordinate], List[str]]:
    pages: List[DocumentPage] = []
    reading_order: List[Dict[str, Any]] = []
    coordinates: List[DocumentCoordinate] = []
    warnings: List[str] = []

    body = doc.element.body
    current_page_number = 1
    page_position = 1
    paragraph_index = 0
    table_index = 0
    page_prefix = build_source_path("document", "page", current_page_number)
    current_page = DocumentPage(
        id=make_node_id("page", current_page_number, page_position),
        page_number=current_page_number,
        position=page_position,
        title=file_name,
        metadata={"file_name": file_name},
    )

    def flush_page() -> None:
        nonlocal current_page, current_page_number, page_position, page_prefix
        pages.append(current_page)
        page_position += 1
        current_page_number += 1
        page_prefix = build_source_path("document", "page", current_page_number)
        current_page = DocumentPage(
            id=make_node_id("page", current_page_number, page_position),
            page_number=current_page_number,
            position=page_position,
            title=file_name,
            metadata={"file_name": file_name},
        )

    for child in body:
        if child.tag.endswith("p"):
            paragraph = Paragraph(child, doc)
            paragraph_model = build_paragraph_models(paragraph, current_page_number, len(current_page.paragraphs) + 1, paragraph_index, page_prefix)
            paragraph_model.is_page_break_before = _has_page_break_before(child)
            paragraph_model.contains_page_break = _has_page_break_within(child)
            current_page.paragraphs.append(paragraph_model)
            current_page.reading_order.append(
                {
                    "order": len(current_page.reading_order) + 1,
                    "type": "paragraph",
                    "id": paragraph_model.id,
                    "text": paragraph_model.text,
                    "source_path": paragraph_model.source_path,
                }
            )
            reading_order.append(
                {
                    "page_number": current_page_number,
                    "order": len(reading_order) + 1,
                    "type": "paragraph",
                    "id": paragraph_model.id,
                    "text": paragraph_model.text,
                    "source_path": paragraph_model.source_path,
                }
            )
            coordinates.append(
                DocumentCoordinate(
                    id=make_coordinate_id("paragraph", paragraph_model.id),
                    object_type="paragraph",
                    source_id=paragraph_model.id,
                    source_path=paragraph_model.source_path or "",
                    page_number=current_page_number,
                    position=paragraph_model.position,
                    text=paragraph_model.text,
                )
            )
            paragraph_index += 1

            if paragraph_model.is_page_break_before and (len(current_page.paragraphs) > 1 or current_page.tables):
                flush_page()
                continue

            if paragraph_model.contains_page_break:
                flush_page()
                continue

        elif child.tag.endswith("tbl"):
            table = Table(child, doc)
            table_model = parse_table(table, current_page_number, len(current_page.tables) + 1, table_index, page_prefix)
            current_page.tables.append(table_model)
            current_page.reading_order.append(
                {
                    "order": len(current_page.reading_order) + 1,
                    "type": "table",
                    "id": table_model.id,
                    "table_number": table_model.table_number,
                    "source_path": table_model.source_path,
                }
            )
            reading_order.append(
                {
                    "page_number": current_page_number,
                    "order": len(reading_order) + 1,
                    "type": "table",
                    "id": table_model.id,
                    "table_number": table_model.table_number,
                    "source_path": table_model.source_path,
                }
            )
            coordinates.append(
                DocumentCoordinate(
                    id=make_coordinate_id("table", table_model.id),
                    object_type="table",
                    source_id=table_model.id,
                    source_path=table_model.source_path or "",
                    page_number=current_page_number,
                    position=table_model.position,
                    table_number=table_model.table_number,
                    text="",
                )
            )
            for row in table_model.rows:
                for cell in row.cells:
                    coordinates.append(
                        DocumentCoordinate(
                            id=make_coordinate_id("cell", cell.id),
                            object_type="cell",
                            source_id=cell.id,
                            source_path=cell.source_path or "",
                            page_number=current_page_number,
                            position=cell.position,
                            table_number=table_model.table_number,
                            row_index=cell.row_index,
                            column_index=cell.column_index,
                            rowspan=cell.rowspan,
                            colspan=cell.colspan,
                            text=cell.text,
                        )
                    )
            table_index += 1

    pages.append(current_page)
    return pages, reading_order, coordinates, warnings

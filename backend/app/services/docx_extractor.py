import io
import re
import logging
from docx import Document
from typing import List

logger = logging.getLogger("docx_extractor")

class DocxExtractionService:
    @staticmethod
    def extract_raw_text(file_bytes: bytes, file_name: str) -> str:
        file_name_lower = file_name.lower() if file_name else ""
        
        if file_name_lower.endswith(".doc"):
            logger.info("Extracting legacy .doc file using binary decoder fallback")
            return DocxExtractionService._extract_doc_fallback(file_bytes)
            
        try:
            logger.info("Extracting modern .docx file")
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            full_text = []
            
            # Extract Paragraphs
            for p in doc.paragraphs:
                full_text.append(p.text)
                
            # Extract Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text or "")
                    full_text.append("\t".join(row_text))
                full_text.append("")
                
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Failed to parse docx using standard parser: {e}. Falling back to binary decoder.")
            return DocxExtractionService._extract_doc_fallback(file_bytes)

    @staticmethod
    def _extract_doc_fallback(file_bytes: bytes) -> str:
        # A legacy .doc file contains text in either UTF-16LE or Latin-1
        try:
            text_utf16 = file_bytes.decode('utf-16le', errors='ignore')
            words_utf16 = re.findall(r'[\u0020-\u007E\u00A0-\u00FF\n\r\t]{4,}', text_utf16)
            utf16_extracted = "\n".join(words_utf16)
        except Exception:
            utf16_extracted = ""

        try:
            text_latin = file_bytes.decode('latin-1', errors='ignore')
            words_latin = re.findall(r'[\x20-\x7E\x0A\x0D\x09]{4,}', text_latin)
            latin_extracted = "\n".join(words_latin)
        except Exception:
            latin_extracted = ""

        return utf16_extracted if len(utf16_extracted) > len(latin_extracted) else latin_extracted

    @staticmethod
    def split_into_chunks(text: str, chunk_size: int = 2000) -> List[str]:
        chunks = []
        if not text or not text.strip():
            return chunks
            
        current_pos = 0
        while current_pos < len(text):
            end = min(current_pos + chunk_size, len(text))
            
            if end < len(text):
                last_newline = text.rfind("\n", current_pos, end)
                if last_newline > current_pos:
                    end = last_newline
                    
            chunks.append(text[current_pos:end].strip())
            current_pos = end
            
            if len(chunks) > 100:
                break
                
        return chunks

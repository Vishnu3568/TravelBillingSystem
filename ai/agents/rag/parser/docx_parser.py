import io
import re
from docx import Document
from .parser_interface import BaseParser

class DocxParser(BaseParser):
    def parse(self, file_bytes: bytes, file_name: str) -> str:
        file_name_lower = file_name.lower() if file_name else ""
        
        # Check if legacy binary .doc format
        if file_name_lower.endswith(".doc"):
            return self._extract_doc_fallback(file_bytes)
            
        try:
            doc = Document(io.BytesIO(file_bytes))
            full_text = []
            
            body = doc.element.body
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            
            for child in body:
                if child.tag.endswith('p'):
                    p = Paragraph(child, doc)
                    if p.text.strip():
                        full_text.append(p.text.strip())
                elif child.tag.endswith('tbl'):
                    t = Table(child, doc)
                    for row in t.rows:
                        row_data = [cell.text.strip() if cell.text else "" for cell in row.cells]
                        cleaned_row = []
                        prev = None
                        for c in row_data:
                            if c == prev and c != "":
                                cleaned_row.append("")
                            else:
                                cleaned_row.append(c)
                                prev = c
                        full_text.append(" | ".join(cleaned_row))
                    full_text.append("")
                    
            return "\n".join(full_text)
        except Exception as e:
            try:
                # Direct simple paragraph fallback
                doc = Document(io.BytesIO(file_bytes))
                return "\n".join([p.text for p in doc.paragraphs])
            except Exception:
                # Return cleaned binary text fallback for corrupt docx or legacy doc
                return self._extract_doc_fallback(file_bytes)

    def _extract_doc_fallback(self, file_bytes: bytes) -> str:
        try:
            text_utf16 = file_bytes.decode('utf-16le', errors='ignore')
            words_utf16 = re.findall(r'[\u0020-\u007E\u00A0-\u00FF\n\r\t]{4,}', text_utf16)
            utf16_extracted = "\n".join(words_utf16)
        except Exception:
            utf16_extracted = ""

        try:
            text_latin = file_bytes.decode('latin-1', errors='ignore')
            words_latin = re.findall(r'[\x20-\x7E\x0A\x0D\x09]{4,}', text_latin)
            latin_extracted = "\n".join(words_latin)
        except Exception:
            latin_extracted = ""

        return utf16_extracted if len(utf16_extracted) > len(latin_extracted) else latin_extracted
